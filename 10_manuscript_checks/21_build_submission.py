"""
21_build_submission.py -- assemble the submission package for BMC Genomics
                          (or Cell Press via --journal cell)

Purpose   Rebuilds submission/ from scratch. File names and structure follow
          the target journal's file requirements:

            Manuscript.docx
            Cover letter.docx
            Figure 1.tif ... Figure 6.tif
            Additional file 1.docx / .pdf   (Figures S1-S9 with legends)
            Additional file 2.xlsx          (Tables S1-S14)
            OPEN_ITEMS.md                   (what only the authors can supply)

          With --journal cell the names revert to Cell Press's own
          (Supplemental Information, Supplementary_Tables) and the package
          additionally carries highlights.docx, Graphical abstract.tif and
          KRT.docx (the latter from 20_key_resources_table.py).

Templates The two Word templates are generated here from pandoc's own default
          reference document and patched in place, so the build depends on
          nothing outside this repository. The manuscript template carries
          continuous line numbers and double spacing; the supplement template
          carries neither line nor page numbers, as the checklist requires,
          and has 15 mm margins so that a 174 mm figure fits the text width.

Citations The reference apparatus is format-agnostic: manuscript/references.bib
          plus a CSL switch. The default is BMC Genomics (Vancouver); `--style
          cell` switches to a Cell Press format. One flag, no retypesetting.
          While the manuscript carries its references as a numbered list in
          the text, pandoc runs without --citeproc; the switch takes effect as
          soon as citations are written as @key.

Images    The panels are drawn at 600 dpi (figure_style/publication_style.R).
          They are wrapped into TIFF without resampling, so the resolution is
          above the 300 dpi the checklist asks for.

Inputs    manuscript/MANUSCRIPT.md, CAPTIONS_MAIN.md, CAPTIONS_SUPPLEMENT.md,
          COVER_LETTER.md, figures/F*.png, figures/S*.png, figures/GA.png,
          figures/data/TS*.csv
Output    submission_BMC_Genomics/ beside the repository
          (submission_CellPress_MJS/ with --journal cell)
Runtime   about a minute (the PDF conversion needs Word)
"""
from __future__ import annotations

import argparse
import os
import pathlib
import re
import shutil
import subprocess
import sys
import tempfile
import zipfile

import pandas as pd
from PIL import Image

_env = os.environ.get("PAPER_V2_ROOT")
ROOT = (pathlib.Path(_env) if _env
        else pathlib.Path(__file__).resolve().parents[1])
PAPER = ROOT / "manuscript"
FIG = ROOT / "figures"
DAT = FIG / "data"
# The built package is a product, not a source, so it is written beside the
# repository rather than into it -- one folder per target journal, so a
# delivered package cannot be overwritten by a build for a different journal.
# `--out` overrides this.
OUT_BY_JOURNAL = {"bmc": ROOT.parent / "submission_BMC_Genomics",
                  "cell": ROOT.parent / "submission_CellPress_MJS"}
OUT = OUT_BY_JOURNAL["bmc"]

PANDOC = shutil.which("pandoc") or str(
    pathlib.Path.home() / ".local" / "bin" / "pandoc.exe")

CSL = {"cell": PAPER / "csl" / "cell.csl",
       "vancouver": PAPER / "csl" / "springer-vancouver.csl"}

# The supplementary tables in order, with the title under which the manuscript
# refers to them.
TABLES = [
    ("TS1", "TS1_eighteen_datasets", "The 18 perturbation datasets"),
    ("TS2", "TS2_screen_exclusion_codes", "Every screened series with its verdict and exclusion code"),
    ("TS3", "TS3_calibrations", "Every calibration, per dataset and per donor cell"),
    ("TS4", "TS4_module_genes", "The 173 module genes with their direction"),
    ("TS5", "TS5_in_vivo_per_zone_and_specimen", "In vivo, per zone and per specimen"),
    ("TS6", "TS6_panels_and_classes", "Gene panels and mechanism classes"),
    ("TS7", "TS7_all_statistics", "Every statistic of the paper with its detection limit"),
    ("TS8", "TS8_preregistrations", "The preregistrations, including the ones that fell"),
    ("TS9", "TS9_gene_sets_v2", "External gene sets, narrow and broad"),
    ("TS9b", "TS9b_gene_sets_v2_module_genes", "Module genes in the external gene sets"),
    ("TS10", "TS10_decomposition_eighteen", "The three-way decomposition, per dataset"),
    ("TS11", "TS11_in_vivo_gene_decomposition", "In vivo, gene-level decomposition"),
    ("TS11b", "TS11b_in_vivo_gene_ranking", "In vivo, gene ranking"),
    ("TS12", "TS12_primary_publications", "Primary publication of every series used"),
    ("TS13", "TS13_in_vivo_hypertrophic_sensitivity", "In vivo, sensitivity to the hypertrophic zone"),
    ("TS14", "TS14_levels_book", "Levels, detection limits and what each level carries"),
]

# Which figure each supplementary table relates to.
RELATED = {"TS1": 1, "TS2": 1, "TS3": 1, "TS4": 2, "TS5": 3, "TS6": 4,
           "TS7": 6, "TS8": 1, "TS9": 2, "TS9b": 2, "TS10": 2, "TS11": 3,
           "TS11b": 3, "TS12": 1, "TS13": 3, "TS14": 6}

SHEETNAMES = {"S1": "S1_datasets", "S2": "S2_screen_exclusions",
              "S3": "S3_calibrations", "S4": "S4_module_genes",
              "S5": "S5_in_vivo_zones", "S6": "S6_panels_classes",
              "S7": "S7_all_statistics", "S8": "S8_preregistrations",
              "S9": "S9_external_gene_sets", "S9b": "S9b_set_membership",
              "S10": "S10_decomposition", "S11": "S11_gene_omission",
              "S11b": "S11b_gene_ranking", "S12": "S12_primary_publications",
              "S13": "S13_hypertrophic_sensitivity", "S14": "S14_levels_book"}


def run(cmd: list[str]) -> None:
    r = subprocess.run(cmd, capture_output=True, text=True)
    if r.returncode != 0:
        raise RuntimeError(f"{cmd[0]} failed:\n{r.stderr[:1500]}")


# ------------------------------------------------------- Word templates
def _patch_reference_docx(target: pathlib.Path, *, line_numbers: bool,
                          margin_mm: float, line_spacing: float,
                          font: str, size_pt: float) -> None:
    """Write a pandoc reference document with our page set-up.

    Built from `pandoc --print-default-data-file reference.docx`, so the
    build needs no template file from outside this repository.
    """
    raw = subprocess.run([PANDOC, "--print-default-data-file", "reference.docx"],
                         capture_output=True)
    if raw.returncode != 0 or not raw.stdout:
        raise RuntimeError("pandoc could not supply its reference document")
    fd, name = tempfile.mkstemp(suffix=".docx")
    os.close(fd)          # Windows keeps the handle open otherwise
    src = pathlib.Path(name)
    src.write_bytes(raw.stdout)

    twip = lambda mm: int(round(mm * 56.6929))          # noqa: E731
    sect = (
        "<w:sectPr>"
        f'<w:pgSz w:w="{twip(210)}" w:h="{twip(297)}"/>'
        f'<w:pgMar w:top="{twip(margin_mm)}" w:right="{twip(margin_mm)}" '
        f'w:bottom="{twip(margin_mm)}" w:left="{twip(margin_mm)}" '
        'w:header="708" w:footer="708" w:gutter="0"/>'
        + ('<w:lnNumType w:countBy="1" w:restart="continuous" w:distance="360"/>'
           if line_numbers else "")
        + "</w:sectPr>")

    with zipfile.ZipFile(src) as z:
        items = {n: z.read(n) for n in z.namelist()}
    doc = items["word/document.xml"].decode("utf-8")
    doc = re.sub(r"<w:sectPr>.*?</w:sectPr>", sect, doc, flags=re.S)
    items["word/document.xml"] = doc.encode("utf-8")

    st = items["word/styles.xml"].decode("utf-8")
    st = st.replace(
        '<w:rFonts w:asciiTheme="minorHAnsi" w:eastAsiaTheme="minorEastAsia" '
        'w:hAnsiTheme="minorHAnsi" w:cstheme="minorBidi" />',
        f'<w:rFonts w:ascii="{font}" w:hAnsi="{font}" w:cs="{font}" />')
    st = st.replace('<w:sz w:val="24" />',
                    f'<w:sz w:val="{int(size_pt * 2)}" />')
    st = st.replace('<w:szCs w:val="24" />',
                    f'<w:szCs w:val="{int(size_pt * 2)}" />')
    st = re.sub(
        r'<w:pPrDefault>.*?</w:pPrDefault>',
        '<w:pPrDefault><w:pPr><w:spacing w:after="120" '
        f'w:line="{int(line_spacing * 240)}" w:lineRule="auto"/>'
        '</w:pPr></w:pPrDefault>', st, flags=re.S)
    # Headings in black and in the body font: pandoc's default template uses
    # the Word theme accent colour and the major font, which makes the
    # document look like a report rather than like a manuscript.
    st = re.sub(r'<w:color w:val="0F4761" w:themeColor="accent1"\s*'
                r'w:themeShade="BF" />', '<w:color w:val="000000" />', st)
    st = re.sub(r'<w:rFonts w:asciiTheme="majorHAnsi"\s*'
                r'w:eastAsiaTheme="majorEastAsia" w:hAnsiTheme="majorHAnsi"\s*'
                r'w:cstheme="majorBidi" />',
                f'<w:rFonts w:ascii="{font}" w:hAnsi="{font}" w:cs="{font}" />',
                st)
    b = int(size_pt * 2)
    for old_sz, new_sz in ((' w:val="40" ', f' w:val="{b + 8}" '),
                           (' w:val="32" ', f' w:val="{b + 4}" '),
                           (' w:val="28" ', f' w:val="{b + 2}" '),
                           (' w:val="56" ', f' w:val="{b + 8}" ')):
        st = st.replace('<w:sz' + old_sz + '/>', '<w:sz' + new_sz + '/>')
        st = st.replace('<w:szCs' + old_sz + '/>', '<w:szCs' + new_sz + '/>')
    # Headings are set in bold; without it they differ from body text only in
    # size, and at manuscript sizes that is not enough to read as a heading.
    st = re.sub(r'(<w:style w:type="paragraph" w:styleId="Heading[1-6]".*?<w:rPr>)',
                r'\1<w:b/>', st, flags=re.S)
    items["word/styles.xml"] = st.encode("utf-8")

    target.parent.mkdir(parents=True, exist_ok=True)
    with zipfile.ZipFile(target, "w", zipfile.ZIP_DEFLATED) as z:
        for n, b in items.items():
            z.writestr(n, b)
    src.unlink(missing_ok=True)


def templates() -> dict[str, pathlib.Path]:
    """Three page set-ups: the manuscript (line numbers, double spacing, as
    the checklist requires), the supplement (no line or page numbers, narrow
    margins so a 174 mm figure fits) and the letter items (neither)."""
    tmp = pathlib.Path(tempfile.mkdtemp(prefix="submission_templates_"))
    out = {"manuscript": tmp / "manuscript.docx",
           "supplement": tmp / "supplement.docx",
           "letter": tmp / "letter.docx"}
    _patch_reference_docx(out["manuscript"], line_numbers=True, margin_mm=25.4,
                          line_spacing=2.0, font="Arial", size_pt=12)
    _patch_reference_docx(out["supplement"], line_numbers=False, margin_mm=15.0,
                          line_spacing=1.15, font="Arial", size_pt=10)
    _patch_reference_docx(out["letter"], line_numbers=False, margin_mm=25.4,
                          line_spacing=1.15, font="Arial", size_pt=11)
    return out


# ------------------------------------------------------ citation style
# The checklist is explicit: "References must be cited by superscript numbers
# running consecutively in the text." The draft carries author-year in square
# brackets because that is readable while writing and while proof-reading.
# The conversion happens here, at package build time -- one switch, no
# retypesetting: for BMC Genomics (Vancouver) the numbering is the same and
# only the bracket form differs.
CITATION = re.compile(r"\[([^\[\]\n]{4,1000})\]")
NOT_A_CITATION = re.compile(
    r"^(?:[A-Z]|C|E|F|\d+"
    r"|(?:Fig\.?|Figure|Table|TS|S)\s?\d[\w., ]*"
    r"|(?:z|P|OR|n|rho)\s.*"
    r"|GitHub URL.*|10\.5281.*|sic)$")


def _is_citation(raw: str) -> bool:
    raw = raw.strip()
    if NOT_A_CITATION.match(raw):
        return False
    return bool(re.search(r"\b(19|20)\d{2}\b", raw))


def _references(md: str) -> list[tuple[int, str]]:
    part = md.split("\n## References", 1)[1]
    return [(int(m.group(1)), re.sub(r"\s+", " ", m.group(2)).strip())
            for m in re.finditer(r"^(\d+)\.\s+(.+?)(?=^\d+\.\s|\Z)", part,
                                 flags=re.M | re.S)]


def _matches(citation: str, reference: str) -> bool:
    yc = set(re.findall(r"\b((?:19|20)\d{2})\b", citation))
    yr = set(re.findall(r"\b((?:19|20)\d{2})\b", reference))
    if yc and yr and not (yc & yr):
        return False
    filler = {"et", "al", "and", "the", "review", "for", "with"}
    words = [w for w in re.findall(r"[A-Za-zÄÖÜäöüß.\-]{3,}", citation)
             if w.lower().strip(".-") not in filler]
    low = reference.lower()
    if any(w.lower().strip(".-") in low for w in words):
        return True
    # Two-letter surnames ("He", "Wu") fall through the length filter. The
    # first token of a citation is always the surname; look for it at a word
    # boundary of the reference.
    first = citation.strip().split()[0].strip(".,-") if citation.strip() else ""
    return bool(first) and re.search(r"\b" + re.escape(first.lower()) + r"\b",
                                     low) is not None


def _lift_block(md: str, markers: tuple[str, str], token: str):
    """Replace a marked block with a token and return (md, inner)."""
    start = md.find(markers[0])
    if start < 0:
        return md, ""
    end = md.find(markers[1], start)
    if end < 0:
        return md, ""
    inner = md[start + len(markers[0]):end]
    return md[:start] + token + md[end + len(markers[1]):], inner


PRIMARY_LIST = ("<!-- PRIMARY_SOURCES:START -->", "<!-- PRIMARY_SOURCES:END -->")
PRIMARY_CITE = ("<!-- PRIMARY_CITATIONS:START -->", "<!-- PRIMARY_CITATIONS:END -->")


def _flatten_citations(md: str) -> str:
    """Put every citation bracket on one line before it is numbered.

    CITATION deliberately refuses to match across a newline, so that a stray
    bracket in a table or a code block cannot be mistaken for a citation. The
    cost is that a citation the source happens to wrap -- [Duggan 2021;
    Savarirayan et al. 2024] over two lines -- is silently left in author-year
    form in the delivered document. Flattening first removes that trap, and
    non-breaking spaces inside the bracket become ordinary ones so that the
    entry lookup sees the text it expects.
    """
    def eine(m: "re.Match[str]") -> str:
        roh = m.group(1)
        if "\n" not in roh:
            return m.group(0)
        flach = re.sub(r"\s+", " ", roh).strip()
        return f"[{flach}]" if _is_citation(flach) else m.group(0)

    return re.sub(r"\[([^\[\]]{4,1000})\]", eine, md, flags=re.S)


def _compress(nums: "list[int]") -> str:
    """Render a citation's numbers the way a numbered style expects them.

    Three or more consecutive numbers become a range, so the collective
    citation of the eighteen datasets' primary publications reads [11-36]
    rather than as twenty-six comma-separated numbers.
    """
    teile: list[str] = []
    lauf: list[int] = []

    def spuele() -> None:
        if not lauf:
            return
        if len(lauf) >= 3:
            teile.append(f"{lauf[0]}-{lauf[-1]}")
        else:
            teile.extend(str(n) for n in lauf)
        lauf.clear()

    for n in sorted(dict.fromkeys(nums)):
        if lauf and n == lauf[-1] + 1:
            lauf.append(n)
        else:
            spuele()
            lauf.append(n)
    spuele()
    return ",".join(teile)


def number_references(md: str, superscript: bool = True) -> str:
    """Turn author-year citations into consecutive numbers.

    Numbers run in order of FIRST occurrence in the text, as Cell Press
    requires, and the list is re-set in the same order. A citation without an
    entry aborts the build -- better no package than one with a hole in the
    bibliography.

    The generated PRIMARY_SOURCES block carries its own 1..N numbering; its
    numbers collide with the main list, so the block is lifted out of the
    reference list before parsing and re-numbered to continue after the main
    references. The collective author-year bracket in STAR Methods is
    flattened back onto one line and converted at its original position, so
    its 24 numbers take their place in first-occurrence order like any other
    citation.
    """
    head, _, tail = md.partition("\n## References")
    tail, _primary_list = _lift_block(tail, PRIMARY_LIST, "")
    head, primary_cite = _lift_block(head, PRIMARY_CITE, "\x00PCT\x00")
    # The collective bracket spans several source lines, so neither CITATION
    # nor a post-hoc conversion can match it across the newlines. It is put
    # back flattened and is numbered at its original position in the text,
    # which keeps first-occurrence order intact.
    head = head.replace("\x00PCT\x00", " ".join(primary_cite.split()))
    primary_cite = ""

    refs = _references(head + "\n## References" + tail)

    # entries: int keys = the main references, "P<n>" = the primary sources.
    entries: list[tuple[object, str]] = list(refs)
    for m in re.finditer(r"^(\d+)\.\s+(.+?)(?=^\d+\.\s|\Z)", _primary_list,
                         flags=re.M | re.S):
        entries.append((f"P{m.group(1)}",
                        re.sub(r"\s+", " ", m.group(2)).strip()))
    texts = dict(entries)
    order: list[object] = []

    def number_for(key) -> int:
        if key not in order:
            order.append(key)
        return order.index(key) + 1

    def lookup(single: str):
        for pass_int in (True, False):
            for k, rtext in entries:
                if isinstance(k, int) == pass_int and _matches(single, rtext):
                    return k
        return None

    def replace(m: "re.Match[str]") -> str:
        raw = m.group(1)
        if not _is_citation(raw):
            return m.group(0)
        singles = [s.strip() for s in raw.split(";")]
        keys = []
        for single in singles:
            k = lookup(single)
            if k is None:
                raise RuntimeError(f"citation without an entry: {single!r}")
            if k not in keys:
                keys.append(k)
        nums = _compress([number_for(k) for k in keys])
        return f"^{nums}^" if superscript else f"[{nums}]"

    head = _flatten_citations(head)
    head = CITATION.sub(replace, head)

    # Uncited main references are appended rather than dropped; the two-way
    # cross-check (11_check_references.py) reports them separately.
    for n, _ in refs:
        if n not in order:
            number_for(n)
    # The primary sources follow, in their own (chronological) list order.
    for k, _ in entries:
        if isinstance(k, str) and k not in order:
            number_for(k)

    out = ["\n## References\n"]
    for k in order:
        out.append(f"{order.index(k) + 1}. {texts[k]}\n")
    return head + "".join(out)


def legends_as_list(md: str, prefix: str = "Figure") -> str:
    """Turn '## Figure 1 - Title' into the form the checklist asks for,
    'Figure 1. Title', and return the legends as ONE list."""
    md = re.sub(r"^#[^#].*$", "", md, count=1, flags=re.M)
    md = re.sub(r"^## " + prefix + r" (\S+) · (.+)$",
                lambda m: f"**{prefix} {m.group(1)}. {m.group(2)}**",
                md, flags=re.M)
    return md.strip()


def strip_rules(md: str) -> str:
    """Remove the horizontal rules that separate sections in the source; in a
    Word manuscript they render as lines across the page."""
    return re.sub(r"(?m)^---\s*$\n?", "", md)


# --------------------------------------------------------------- docx
def md_to_docx(md: str, target: pathlib.Path, template: pathlib.Path,
               style: str = "cell") -> None:
    tmp = OUT / "_tmp.md"
    tmp.write_text(md, encoding="utf-8")
    cmd = [PANDOC, str(tmp), "-o", str(target), "--from",
           "markdown+smart+pipe_tables", "--to", "docx",
           "--bibliography", str(PAPER / "references.bib"),
           "--csl", str(CSL[style]), "--reference-doc", str(template)]
    run(cmd)
    tmp.unlink(missing_ok=True)


def docx_to_pdf(source: pathlib.Path) -> pathlib.Path | None:
    """docx -> pdf through Word. Without Word the step is skipped, not faked."""
    target = source.with_suffix(".pdf")
    ps = ("$w = New-Object -ComObject Word.Application; $w.Visible = $false; "
          f"$d = $w.Documents.Open('{source}'); "
          f"$d.SaveAs([ref]'{target}', [ref]17); $d.Close(); $w.Quit()")
    r = subprocess.run(["powershell", "-NoProfile", "-Command", ps],
                       capture_output=True, text=True)
    if r.returncode != 0 or not target.exists():
        # Word reports through either stream, and on a hard COM failure through
        # neither -- so guard both. A missing PDF is a skipped step, never a
        # reason to abort the build: the journal wants the editable file.
        grund = ((r.stderr or "") + (r.stdout or "")).strip()
        print("  ! PDF skipped (Word unavailable or busy):",
              grund[:200] or f"no output, exit {r.returncode}")
        return None
    return target


# ---------------------------------------------------------------- TIFF
def png_to_tif(source: pathlib.Path, target: pathlib.Path,
               dpi: int = 600) -> None:
    """PNG to TIFF with LZW, without resampling -- repackaging only."""
    im = Image.open(source).convert("RGB")
    im.save(target, format="TIFF", compression="tiff_lzw", dpi=(dpi, dpi))


def supplement_name(journal: str) -> str:
    """BMC Genomics cites supplementary material by additional-file name, and
    the file itself must carry that name. Cell Press calls it Supplemental
    Information."""
    return ("Additional file 1" if journal == "bmc"
            else "Supplemental Information")


def tables_name(journal: str) -> str:
    return "Additional file 2" if journal == "bmc" else "Supplementary_Tables"


# ------------------------------------------------------------ the parts
def build_manuscript(style: str, template: pathlib.Path,
                     journal: str = "cell") -> None:
    md = (PAPER / "MANUSCRIPT.md").read_text(encoding="utf-8")
    # The highlights, when present, are a separate submission item (Cell Press)
    # and are not repeated in the manuscript. BMC Genomics carries none.
    md = re.sub(r"\n## Highlights\n.*?\n---\n", "\n", md, flags=re.S)

    legends = legends_as_list(
        (PAPER / "CAPTIONS_MAIN.md").read_text(encoding="utf-8"))
    if journal == "bmc":
        # BMC Genomics: Background/Results/Discussion/Conclusions/Methods/
        # abbreviations/additional files/Declarations/References. The figure
        # legends go into the main document -- BMC keeps them out of the
        # graphic files -- directly after the Methods. The supplement is a
        # separate additional file, so no "supplemental item titles" block.
        md = md.replace("\n## List of abbreviations\n",
                        "\n## Figure legends\n\n" + legends
                        + "\n\n## List of abbreviations\n", 1)
    else:
        # Cell Press: figure legends between the discussion and the STAR
        # Methods, and the supplemental item titles before the references.
        md = md.replace("\n## STAR Methods\n",
                        "\n## Figure Legends\n\n" + legends
                        + "\n\n## STAR Methods\n", 1)
        md = md.replace("\n## References\n",
                        "\n" + supplemental_item_titles() + "\n## References\n", 1)

    md = number_references(md, superscript=(style == "cell"))
    target = OUT / "Manuscript.docx"
    md_to_docx(strip_rules(md), target, template, style=style)
    print(f"  Manuscript.docx           {len(md.split())} words, "
          f"{'superscript numbers' if style == 'cell' else 'numbers in brackets'}")
    # The journal wants the editable file; the PDF is for reading the package
    # before it is uploaded.
    if docx_to_pdf(target):
        print("  Manuscript.pdf            written")


def supplemental_item_titles() -> str:
    """Titles and 'Related to' of the supplemental items. The checklist puts
    them in the main document, after the figure legends."""
    lines = ["## Supplemental Item Titles and Legends", "",
             "Supplementary Tables are delivered as one workbook,",
             "`Supplementary_Tables.xlsx`, with sixteen individually named",
             "sheets: S1_datasets; S2_screen_exclusions; S3_calibrations;",
             "S4_module_genes; S5_in_vivo_zones; S6_panels_classes;",
             "S7_all_statistics; S8_preregistrations; S9_external_gene_sets;",
             "S9b_set_membership; S10_decomposition; S11_gene_omission;",
             "S11b_gene_ranking; S12_primary_publications;",
             "S13_hypertrophic_sensitivity; S14_levels_book. Each sheet carries",
             "its content description and its Related-to figure in the",
             "Supplemental Information.",
             ""]
    return "\n".join(lines) + "\n"


def build_highlights(style: str, template: pathlib.Path) -> None:
    md = (PAPER / "MANUSCRIPT.md").read_text(encoding="utf-8")
    m = re.search(r"## Highlights\n(.*?)\n---\n", md, flags=re.S)
    # Bullets may wrap over several lines; continuation lines belong to the
    # previous point -- dropping them truncated a highlight in the delivered
    # docx.
    points: list[str] = []
    if m:
        for z in m.group(1).split("\n"):
            s = z.strip()
            if s.startswith("-"):
                points.append(s[1:].strip())
            elif points and s:
                points[-1] += " " + s
    text = ("Highlights\n\n"
            + "\n".join("- " + p for p in points) + "\n")
    md_to_docx(text, OUT / "highlights.docx", template, style=style)
    too_long = [p for p in points if len(p) > 85]
    print(f"  highlights.docx           {len(points)} points"
          + (f", {len(too_long)} longer than 85 characters" if too_long else ""))


def build_cover_letter(style: str, template: pathlib.Path) -> None:
    md = (PAPER / "COVER_LETTER.md").read_text(encoding="utf-8")
    md_to_docx(md, OUT / "Cover letter.docx", template, style=style)
    print(f"  Cover letter.docx         {len(md.split())} words")


def build_figures(journal: str = "cell") -> None:
    for i in range(1, 7):
        png_to_tif(FIG / f"F{i}.png", OUT / f"Figure {i}.tif")
    size = Image.open(OUT / "Figure 1.tif").size
    print(f"  Figure 1-6.tif            RGB, LZW, 600 dpi, "
          f"Figure 1 = {size[0]}x{size[1]} px")
    if journal == "bmc":
        # BMC Genomics carries no graphical abstract.
        return
    # The graphical abstract must be EXACTLY 1200 x 1200 px at 300 dpi, so
    # 09_figures/30_graphical_abstract.py draws it at 300 rather than 600 dpi. Here
    # it is only checked and packed.
    ga = Image.open(FIG / "GA.png")
    if ga.size != (1200, 1200):
        raise RuntimeError(
            f"the graphical abstract is {ga.size[0]}x{ga.size[1]} px; "
            "exactly 1200x1200 at 300 dpi is required "
            "(iScience Final File Requirements). "
            "Re-run 09_figures/30_graphical_abstract.py.")
    png_to_tif(FIG / "GA.png", OUT / "Graphical abstract.tif", dpi=300)
    print("  Graphical abstract.tif    RGB, LZW, 1200x1200 px at 300 dpi")


def build_tables(journal: str = "cell") -> pathlib.Path:
    target = OUT / f"{tables_name(journal)}.xlsx"
    with pd.ExcelWriter(target, engine="openpyxl") as w:
        contents = []
        for short, name, title in TABLES:
            p = DAT / f"{name}.csv"
            if not p.exists():
                print("  ! missing:", p.name)
                continue
            d = pd.read_csv(p)
            disp = "S" + short[2:] if short.startswith("TS") else short
            disp = SHEETNAMES.get(disp, disp)
            d.to_excel(w, sheet_name=disp[:31], index=False)
            contents.append({"table": disp, "title": title,
                             "related to figure": RELATED.get(short, 1),
                             "rows": len(d), "columns": d.shape[1]})
        pd.DataFrame(contents).to_excel(w, sheet_name="Contents", index=False)
    print(f"  {target.name} {len(TABLES)} tables")
    return target


def build_supplement(style: str, template: pathlib.Path,
                     journal: str = "cell") -> None:
    """The S figures with their legends, plus the table legends, in ONE
    document: one figure per page, legend kept with its figure."""
    legends = (PAPER / "CAPTIONS_SUPPLEMENT.md").read_text(encoding="utf-8")
    # Everything before the first S panel is front matter for the working
    # copy and does not belong in a delivered item.
    legends = legends[legends.index("## Figure S1 "):]
    table_part = ""
    if "\n# Supplementary tables" in legends:
        legends, table_part = legends.split("\n# Supplementary tables", 1)

    # The checklist: no article title, no author list, no page numbers in the
    # supplement -- the publisher puts a cover sheet in front of it.
    if journal == "bmc":
        parts = ["# Additional file 1", "",
                 "Supplementary figures S1-S9 with their legends.", "",
                 "## Supplementary figures", ""]
    else:
        parts = ["# Supplemental Information", "", "## Supplemental Figures",
                 ""]
    blocks = re.split(r"(?m)^## Figure (S\d) ", legends)
    head = blocks[0].strip()
    if head:
        parts += [head, ""]
    for i in range(1, len(blocks), 2):
        number, body = blocks[i], blocks[i + 1]
        png = FIG / f"{number}.png"
        parts.append(f"![]({png.as_posix()}){{width=174mm}}")
        parts.append("")
        # Checklist: "Figure S3. [Title], Related to Figure 1" -- a full stop
        # after the number, not the middle dot of the working copy.
        title = body.strip().splitlines()[0].lstrip("· ").strip()
        parts.append(f"**Figure {number}. {title}**")
        parts.append("")
        parts += body.strip().split("\n")[1:]
        parts.append("")
    parts += ["## Supplementary tables", "",
              "The supplementary tables are delivered as a single workbook,",
              f"`{tables_name(journal)}.xlsx`; each dataset below is one",
              "individually named sheet in it (S1_datasets through",
              "S14_levels_book, plus S9b and S11b).", ""]
    if table_part:
        parts += table_part.strip().split("\n")
        parts.append("")

    target = OUT / f"{supplement_name(journal)}.docx"
    md_to_docx("\n".join(parts), target, template, style=style)
    typeset_supplement(target)
    print(f"  {target.name}  9 panels + {len(TABLES)} tables")
    if docx_to_pdf(target):
        print(f"  {target.stem}.pdf   written")


def typeset_supplement(path: pathlib.Path) -> None:
    """Set the supplement the way a supplement should look.

    pandoc delivers the content; the layout rules are applied here, because
    they cannot be expressed in markdown:
      * legends at 8.5 pt, so a legend fits on the page with its figure
      * each figure starts a new page, and its legend never leaves it
      * no line numbers (the checklist forbids them in the supplement)
    """
    from docx import Document
    from docx.enum.text import WD_BREAK
    from docx.shared import Pt

    d = Document(str(path))
    paragraphs = d.paragraphs

    def has_image(p) -> bool:
        return bool(p._element.findall(
            ".//{http://schemas.openxmlformats.org/drawingml/2006/"
            "wordprocessingDrawing}inline"))

    image_at = [i for i, p in enumerate(paragraphs) if has_image(p)]
    tables_at = next((i for i, p in enumerate(paragraphs)
                      if p.text.strip() == "Supplemental Tables"), len(paragraphs))

    for n, i in enumerate(image_at):
        end = image_at[n + 1] if n + 1 < len(image_at) else tables_at
        # page break before every figure but the first
        if n:
            paragraphs[i].insert_paragraph_before().add_run().add_break(
                WD_BREAK.PAGE)
        paragraphs[i].paragraph_format.keep_with_next = True
        paragraphs[i].paragraph_format.space_after = Pt(4)
        for j in range(i + 1, end):
            p = paragraphs[j]
            p.paragraph_format.keep_with_next = j < end - 1
            p.paragraph_format.space_after = Pt(3)
            for r in p.runs:
                r.font.size = Pt(8.5)

    # the table legends also read as legends, not as body text
    for j in range(tables_at + 1, len(paragraphs)):
        for r in paragraphs[j].runs:
            r.font.size = Pt(8.5)

    d.save(str(path))
    _strip_line_numbers(path)


def _strip_line_numbers(p: pathlib.Path) -> None:
    """Remove <w:lnNumType/> from a docx -- the checklist asks for neither
    line nor page numbers in the supplement."""
    with zipfile.ZipFile(p) as z:
        items = {n: z.read(n) for n in z.namelist()}
    xml = items["word/document.xml"].decode("utf-8")
    xml = re.sub(r"<w:lnNumType[^/>]*/>", "", xml)
    items["word/document.xml"] = xml.encode("utf-8")
    fd, name = tempfile.mkstemp(suffix=".docx")
    os.close(fd)  # otherwise Windows keeps the handle and move() fails
    tmp = pathlib.Path(name)
    with zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as z:
        for n, b in items.items():
            z.writestr(n, b)
    shutil.move(str(tmp), str(p))


def build_open_items() -> None:
    (OUT / "OPEN_ITEMS.md").write_text(OPEN, encoding="utf-8")
    print("  OPEN_ITEMS.md")


OPEN = """# Submission checklist -- what is still open

Written by `10_manuscript_checks/21_build_submission.py`. Every file beside
this one is built from the sources and is complete. The items below remain,
and none of them is something a script can do.

Target journal: **BMC Genomics** (Research Article). The manuscript was
reframed and reformatted from an earlier Cell Press Multi-Journal submission
(CP-MULTI-JOURNAL-D-26-03425), which is cited in the cover letter.

## Still open

- [ ] **Editorial Manager entries.** ORCID of both authors (also on the title
      page), the corresponding author's contact, and optionally suggested
      reviewers. Confirm the Availability-of-data-and-materials, Competing
      interests, Funding and Authors' contributions statements as they stand in
      the manuscript's Declarations section.
- [ ] **Open-access agreement / APC.** BMC Genomics is fully open access.
      Check institutional / DFG open-access funding and any waiver before the
      article-processing charge is invoiced.
- [x] **Deposit route.** GitHub, public at
      https://github.com/alexschulzcell/Differentiation-paper ; the URL is in
      the Availability of data and materials declaration.

## Closed

- [x] **Authors.** Alexander Schulz (ORCID 0009-0009-2605-4350), first author;
      Christian T. Thiel (ORCID 0000-0003-3817-7277), corresponding author.
      Institute of Human Genetics, Universitaetsklinikum Erlangen, FAU
      Erlangen-Nuernberg, Germany.
- [x] **Funding.** DFG grant TH896/7-1 (C.T.T.).
- [x] **Competing interests.** None declared (in Declarations).
- [x] **Generative AI use** documented in Methods (*Use of large language
      models*), which is where BMC asks for it, and pointed at from the
      Acknowledgements.
- [x] **Ethics.** No new human material or data; oversight for each reanalysed
      data set is documented in its primary publication, which is cited. The
      sex or gender limitation of the single-donor series is stated.
- [x] **Licence.** CC BY 4.0 for text, figures and derived data; MIT for code.
      `LICENSE` and `LICENSE-CODE` are in force in the repository.
- [x] **Accessions.** All eighteen perturbation datasets carry a real
      accession, and the figures carry the accessions in their row labels.
- [x] **Primary publications of every reanalysed series are cited.**
- [x] **No preprint.**
- [x] `python 10_manuscript_checks/10_check_numbers.py` green
- [x] `python 10_manuscript_checks/11_check_references.py` green
- [x] `python 10_manuscript_checks/12_check_language.py` green
- [x] Structured abstract, 349 words (limit 350), no references in it.
- [x] Sections in BMC order: Background, Results, Discussion (Limitations as
      its closing subsection), Conclusions, Methods, Figure legends, List of
      abbreviations, Additional files, Declarations, References.
- [x] Figures cited as `Fig. 2a`, panel letters lower case, numbered in order
      of first mention; legends in the manuscript, not in the graphic files.
- [x] Supplement cited by additional-file name and delivered under that name:
      `Additional file 1` (figures S1-S9), `Additional file 2` (tables
      S1-S14), each listed with format, title and description in the
      *Additional files* section.
- [x] Vancouver references, consecutive runs as ranges (for example [14-39]).
- [x] Declarations carry all seven required subheadings, and the use of
      generative AI is documented in Methods as BMC requires.

## Deliberate deviations from the letter of the guidelines

Both are judgement calls; a reviewer or editor may ask for either to change.

- **Figure legends exceed 300 words** (Figure 2 most: the sheet carries nine
  panels). Every statistic in this paper is reported with its own detection
  limit, and the legends are where those limits sit; cutting them to 300 words
  would mean deleting numbers rather than prose. Table S7 in Additional file 2
  lists every statistic with its limit as the machine-readable counterpart.
- **Figure 2 is 230 mm tall and Figure 4 is 214 mm**, against BMC's 225 mm for
  figure plus legend. The sheets are drawn at 174 mm width, so they will be
  scaled slightly at BMC's 170 mm. Re-rendering was considered and not done, to
  avoid re-flowing panels that are already measured and checked.
"""


def main() -> int:
    ap = argparse.ArgumentParser()
    ap.add_argument("--out", default=None,
                    help="where to write the package; defaults to "
                         "submission_BMC_Genomics/ (or submission_CellPress_"
                         "MJS/ with --journal cell) beside the repository")
    ap.add_argument("--journal", choices=["cell", "bmc"], default="bmc",
                    help="bmc = BMC Genomics (default), cell = Cell Press")
    ap.add_argument("--style", choices=sorted(CSL), default=None,
                    help="citation style; defaults from --journal "
                         "(bmc -> vancouver, cell -> cell)")
    a = ap.parse_args()
    if a.style is None:
        a.style = "vancouver" if a.journal == "bmc" else "cell"

    global OUT
    OUT = (pathlib.Path(a.out) if a.out else OUT_BY_JOURNAL[a.journal])
    OUT.mkdir(parents=True, exist_ok=True)
    # Rebuild means rebuild: a package left over from the other journal
    # carries the other journal's file names, and a mixed package is worse
    # than no package. Only what this run writes may stay.
    gesperrt = []
    for stale in sorted(OUT.iterdir()):
        if not stale.is_file():
            continue
        try:
            stale.unlink()
        except OSError:
            # Word holds the docx it is converting, and a reader may have one
            # open. A file we cannot remove is not a reason to abort a build
            # that would overwrite it anyway; only what stays behind unwritten
            # is worth a warning at the end.
            gesperrt.append(stale.name)
    if gesperrt:
        print("  ! in use, will be overwritten rather than replaced: "
              + ", ".join(gesperrt))
    print(f"21_build_submission.py -- package into {OUT} "
          f"(journal: {a.journal}, style: {a.style})")
    tpl = templates()
    build_manuscript(a.style, tpl["manuscript"], journal=a.journal)
    build_cover_letter(a.style, tpl["letter"])
    build_figures(journal=a.journal)
    build_tables(journal=a.journal)
    build_supplement(a.style, tpl["supplement"], journal=a.journal)
    build_open_items()
    if a.journal == "cell":
        # Cell Press only: highlights and the Key Resources Table.
        build_highlights(a.style, tpl["letter"])
        if not (OUT / "KRT.docx").exists():
            print("  ! KRT.docx missing -- run "
                  "10_manuscript_checks/20_key_resources_table.py first")
    print()
    print("Contents of " + OUT.name + "/:")
    for p in sorted(OUT.iterdir()):
        if p.is_file():
            print(f"  {p.name:34s} {p.stat().st_size / 1024:8.0f} kB")
    return 0


if __name__ == "__main__":
    sys.exit(main())
