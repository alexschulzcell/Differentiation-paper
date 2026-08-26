"""
63_key_resources_table.py -- die Key Resources Table fuer iScience

Zweck    Erzeugt submission/KRT.docx nach dem Aufbau der Vorlage
         `alex phd/Manuscripts/LAMA5 paper/iScience final/KRT.docx`:
         eine dreispaltige Tabelle (REAGENT or RESOURCE | SOURCE | IDENTIFIER)
         mit den Cell-Press-Standardueberschriften in der vorgeschriebenen
         Reihenfolge, jede Zeile einzeilig, keine verbundenen Zellen, keine
         eigene Literaturliste.

         Fuer eine rein rechnerische Arbeit fuellen sich nur drei Rubriken:
         "Deposited data" (die oeffentlichen Datensaetze aus TS1 und die
         beiden Atlanten), "Software and algorithms" (requirements.txt,
         r_packages.txt, die sessionInfo-Dateien) und "Other" (die abgelegten
         Zwischenstaende). Alle uebrigen Rubriken stehen mit N/A da -- die
         Vorlage laesst keine weg.

Eingaben figures/data/TS1_eighteen_datasets.csv
         requirements.txt, r_packages.txt, results/*_sessionInfo.txt
Ausgabe  submission/KRT.docx
Laufzeit Sekunden

Regel 0: der Inhalt ist restlos englisch. Deutsch steht nur in diesen
Kommentaren.
"""
from __future__ import annotations

import os
import pathlib
import re

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

_env = os.environ.get("PAPER_V2_ROOT")
WURZEL = (pathlib.Path(_env) if _env
          else pathlib.Path(__file__).resolve().parents[1])
AUS = WURZEL / "submission"
AUS.mkdir(parents=True, exist_ok=True)

# Die Cell-Press-Rubriken in der vorgeschriebenen Reihenfolge (aus der Vorlage).
RUBRIKEN = [
    "Antibodies",
    "Bacterial and virus strains",
    "Biological samples",
    "Chemicals, peptides, and recombinant proteins",
    "Critical commercial assays",
    "Deposited data",
    "Experimental models: Cell Lines",
    "Experimental models: Organisms/strains",
    "Oligonucleotides",
    "Recombinant DNA",
    "Software and algorithms",
    "Other",
]

# Datensaetze, die nicht in TS1 stehen, aber benutzt werden.
WEITERE_DATEN = [
    ("Human fetal limb single-cell atlas (re-used)", "ArrayExpress",
     "E-MTAB-8813"),
    ("Chromatin accessibility, ATAC-seq (re-used)",
     "Gene Expression Omnibus", "GSE332758"),
    ("Chromatin accessibility, ATAC-seq, second cohort (re-used)",
     "Gene Expression Omnibus", "GSE151311"),
    ("Histone H3K27ac ChIP-seq, second cohort (re-used)",
     "Gene Expression Omnibus", "GSE151315"),
    ("Histone H3K27ac ChIP-seq (re-used)", "Gene Expression Omnibus",
     "GSE129031"),
    ("Promoter methylome, 27K array (re-used)", "Gene Expression Omnibus",
     "GSE33896"),
    ("Promoter methylome, 450K array (re-used)", "Gene Expression Omnibus",
     "GSE129266"),
    ("Postnatal growth plate, single cell (re-used, not calibratable)",
     "Gene Expression Omnibus", "GSE288028"),
    ("Gene constraint (LOEUF)", "gnomAD v4", "RRID:SCR_014964"),
    ("Skeletal dysplasia gene panel (309 green genes)",
     "Genomics England PanelApp", "RRID:SCR_017021"),
    ("Nosology of Genetic Skeletal Disorders (core and broad panels)",
     "Unger et al. 2023", "N/A"),
    ("Height GWAS gene set", "Yengo et al. 2022", "N/A"),
    ("Publication counts per gene", "NCBI gene2pubmed, retrieved 2026-08-22",
     "RRID:SCR_002473"),
    ("Gene annotation", "GENCODE v46 (hg38)", "RRID:SCR_014966"),
]

RRID = {
    "python": "RRID:SCR_008394", "numpy": "RRID:SCR_008633",
    "pandas": "RRID:SCR_018214", "scipy": "RRID:SCR_008058",
    "scanpy": "RRID:SCR_018139", "anndata": "N/A", "h5py": "N/A",
    "leidenalg": "N/A", "igraph": "RRID:SCR_019225",
    "R": "RRID:SCR_001905", "ggplot2": "RRID:SCR_014601",
    "patchwork": "N/A", "ragg": "N/A", "systemfonts": "N/A",
    "DESeq2": "RRID:SCR_015687", "matrixStats": "N/A",
    "rtracklayer": "RRID:SCR_021325", "GenomicRanges": "RRID:SCR_000025",
}


def _pakete() -> list[tuple[str, str, str]]:
    """Software-Zeilen aus requirements.txt und r_packages.txt."""
    zeilen: list[tuple[str, str, str]] = []
    zeilen.append(("Original code (Python and R analysis scripts)",
                   "This paper",
                   "GitHub: [repository URL]"))
    zeilen.append(("Python", "Python Software Foundation",
                   "v3.12.10; " + RRID["python"]))
    for roh in (WURZEL / "requirements.txt").read_text(
            encoding="utf-8").splitlines():
        roh = roh.strip()
        if not roh or roh.startswith("#"):
            continue
        name, _, version = roh.partition("==")
        zeilen.append((name, "PyPI",
                       f"v{version}; {RRID.get(name, 'N/A')}"))
    zeilen.append(("R", "R Foundation", "v4.4.3; " + RRID["R"]))
    for roh in (WURZEL / "r_packages.txt").read_text(
            encoding="utf-8").splitlines():
        roh = roh.strip()
        if not roh or roh.startswith("#"):
            continue
        m = re.match(r"([A-Za-z0-9._]+)_([0-9.\-]+)$", roh)
        if not m:
            continue
        name, version = m.groups()
        zeilen.append((name, "CRAN / Bioconductor",
                       f"v{version}; {RRID.get(name, 'N/A')}"))
    return zeilen


def _daten() -> list[tuple[str, str, str]]:
    ts1 = pd.read_csv(WURZEL / "figures" / "data"
                      / "TS1_eighteen_datasets.csv")
    zeilen: list[tuple[str, str, str]] = []
    gesehen: set[str] = set()
    for _, r in ts1.iterrows():
        gse = str(r["accession"])
        if gse in gesehen:
            continue
        gesehen.add(gse)
        # Alle achtzehn Datensaetze sind oeffentlich; die eigene USC-Serie
        # liegt bei ArrayExpress, alle uebrigen bei GEO.
        archiv = ("ArrayExpress" if gse.startswith("E-MTAB")
                  else "Gene Expression Omnibus")
        zeilen.append((f"Perturbation dataset: {r['dataset']}",
                       archiv, gse))
    zeilen.extend(WEITERE_DATEN)
    return zeilen


def _tabelle() -> dict[str, list[tuple[str, str, str]]]:
    t = {k: [("N/A", "N/A", "N/A")] for k in RUBRIKEN}
    t["Deposited data"] = _daten()
    t["Software and algorithms"] = _pakete()
    t["Other"] = [
        ("Panel-level source data for every figure (CSV)", "This paper",
         "GitHub: [repository URL]; figures/data/"),
        ("Self-test reproducing every number in the manuscript",
         "This paper", "code/70_check_numbers.py"),
        ("Preregistrations, unchanged, including the ones that fell",
         "This paper", "preregistrations/; Supplementary Table 8"),
    ]
    return t


def main() -> None:
    d = Document()
    st = d.styles["Normal"]
    st.font.name = "Arial"
    st.font.size = Pt(9)

    kopf = d.add_paragraph("KEY RESOURCES TABLE")
    kopf.runs[0].bold = True
    kopf.alignment = WD_ALIGN_PARAGRAPH.LEFT

    inhalt = _tabelle()
    n = 1 + sum(1 + len(v) for v in inhalt.values())
    tab = d.add_table(rows=n, cols=3)
    tab.style = "Table Grid"

    for j, h in enumerate(("REAGENT or RESOURCE", "SOURCE", "IDENTIFIER")):
        z = tab.cell(0, j).paragraphs[0]
        z.add_run(h).bold = True

    i = 1
    for rubrik in RUBRIKEN:
        tab.cell(i, 0).paragraphs[0].add_run(rubrik).bold = True
        tab.cell(i, 1).text = ""
        tab.cell(i, 2).text = ""
        i += 1
        for zeile in inhalt[rubrik]:
            for j, wert in enumerate(zeile):
                tab.cell(i, j).text = str(wert)
            i += 1

    p = AUS / "KRT.docx"
    d.save(p)
    print(f"KRT -> {p}  ({n} Zeilen, {len(RUBRIKEN)} Rubriken)")


if __name__ == "__main__":
    main()
